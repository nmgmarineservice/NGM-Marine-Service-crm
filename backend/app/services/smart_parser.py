from docx import Document
import io
import re
import pdfplumber
import pandas as pd
import traceback
import os
import tempfile
import subprocess
import shutil
# Only works on Windows with Word installed
try:
    import win32com.client as win32
    import pythoncom
    win32_error = None
except ImportError as e:
    win32 = None
    win32_error = str(e)

# Environment verification complete
from pptx import Presentation
from typing import List, Dict, Any

class BaseParser:
    @staticmethod
    def map_to_form_fields(structured_data: dict):
        """
        Maps the raw extracted data into the system's FormField schema.
        Includes improved heuristic for:
        1. Grid alignment validation
        2. Yes/No/NA checklist detection
        3. Key-Value pair extraction from 2-column tables
        """
        fields = []
        seen_labels = set()
        
        # 1. Map metadata
        sorted_keys = sorted(structured_data["metadata"].keys())
        for key in sorted_keys:
            value = structured_data["metadata"][key]
            clean_key = re.sub(r'[:\W_]+$', '', key).strip()
            if len(clean_key) < 2: continue
            if clean_key.lower() in seen_labels: continue
            
            # Skip noise
            if clean_key.upper() in ["PAGE", "NMG", "MARINE", "NOON CHIT"]: continue
            
            seen_labels.add(clean_key.lower())

            field_type = "text"
            lk = clean_key.lower()
            if "date" in lk: field_type = "date"
            
            safe_key = re.sub(r'\W+', '_', clean_key).lower()
            fields.append({
                "id": f"f_{safe_key}_{len(fields)}",
                "label": clean_key,
                "type": field_type,
                "required": False,
                "default_value": value
            })

        # 2. Map tables - with Layout-to-Metadata promotion
        for i, table in enumerate(structured_data.get("tables", [])):
            if not table or len(table) < 1: continue
            
            # Prune generic empty/unnamed columns first
            header_row = table[0]
            valid_col_indices = []
            for col_idx in range(len(header_row)):
                col_data = [str(row[col_idx]).strip() for row in table[1:] if col_idx < len(row)]
                has_val = any(c and "Unnamed" not in c for c in col_data)
                h_val = str(header_row[col_idx]).strip()
                if has_val or (h_val and "Unnamed" not in h_val and ":" not in h_val):
                    valid_col_indices.append(col_idx)

            if not valid_col_indices: continue
            
            pruned_table = []
            for row in table:
                pruned_table.append([row[idx] for idx in valid_col_indices if idx < len(row)])
            
            if not pruned_table: continue
            
            headers = [str(h).strip() for h in pruned_table[0]]
            
            # RE-CHECK: Is this just a metadata grid? (e.g. Vessel Name:, Date:, etc.)
            # A metadata grid has ':' in almost every header cell
            is_metadata_grid = all(":" in h or not h for h in headers)
            if is_metadata_grid and len(pruned_table) <= 3:
                for row in pruned_table:
                    for h_val in row:
                        txt = str(h_val).strip()
                        kv = re.search(r'^([^:]{2,50}):\s*(.*)$', txt)
                        if kv:
                            k = kv.group(1).strip()
                            v = kv.group(2).strip()
                            ck = re.sub(r'[:\W_]+$', '', k).strip()
                            if ck.lower() not in seen_labels:
                                seen_labels.add(ck.lower())
                                fields.append({
                                    "id": f"f_meta_{len(fields)}",
                                    "label": ck,
                                    "type": "date" if "date" in ck.lower() else "text",
                                    "default_value": v
                                })
                continue
            
            # 2b. Heuristic: 2-Column Table = Key/Value Metadata (even without colons)
            # If table is exactly 2 columns and has many rows, it's likely a property list
            if len(headers) == 2 and len(pruned_table) > 1:
                is_kv_like = True
                possible_fields = []
                for row_idx, row in enumerate(pruned_table):
                    if len(row) < 2: continue
                    k, v = str(row[0]).strip(), str(row[1]).strip()
                    # If key is too long, probably a paragraph, not a label
                    if len(k) > 60 or len(k) < 2: 
                        is_kv_like = False; break
                    
                    clean_key = re.sub(r'[:\W_]+$', '', k).strip()
                    if clean_key:
                        possible_fields.append((clean_key, v))
                
                if is_kv_like and possible_fields:
                    for k, v in possible_fields:
                        if k.lower() not in seen_labels:
                            seen_labels.add(k.lower())
                            fields.append({
                                "id": f"f_kv_{len(fields)}",
                                "label": k,
                                "type": "date" if "date" in k.lower() else "text",
                                "default_value": v
                            })
                    continue

            # Standard Table Labels
            headers = [h if (h and "Unnamed" not in h) else f"Field {j+1}" for j, h in enumerate(headers)]

            # Specialized: Checklist / Choice Grids (Yes/No, 1-5, etc.)
            # Expanded regex to capture Yes, No, Y, N, Good, Bad, etc.
            rating_pattern = r'^(0[1-5]|[1-5]|Yes|No|Y|N|Good|Bad|Poor|Fair|Sat|Unsat|N\/A|NA)$'
            is_rating_table = any(re.match(rating_pattern, h, re.IGNORECASE) for h in headers)

            if is_rating_table:
                rating_options = [h for h in headers if re.match(rating_pattern, h, re.IGNORECASE)]
                # If we found options but they are mixed with "Question", filter "Question" out
                rating_options = [opt for opt in rating_options if len(opt) < 10] 

                for row in pruned_table[1:]:
                    if not row or not str(row[0]).strip(): continue
                    label = re.sub(r'^\d+[\.\)\s]+', '', str(row[0]).strip())
                    if len(label) > 1:
                        fields.append({
                            "id": f"rating_{i}_{len(fields)}",
                            "label": label, 
                            "type": "select", 
                            "options": rating_options, 
                            "required": False
                        })
                continue

            # Noon Chit / Log Table
            columns = []
            for j, h in enumerate(headers):
                columns.append({"id": f"col_{i}_{j}", "label": h, "type": "text"})

            section_title = f"Form Table {i+1}"
            if any(x in str(headers).upper() for x in ["ITEM", "S.NO", "S. NO.", "PARTICULARS"]):
                section_title = "Data Entry Table"

            fields.append({
                "id": f"table_{i}_{len(fields)}",
                "label": section_title,
                "type": "table",
                "columns": columns
            })

        return fields

class ExcelParser:
    @staticmethod
    def extract(file_content: bytes, filename: str):
        data = {"paragraphs": [], "tables": [], "metadata": {}, "raw_text": ""}
        try:
            engine = 'openpyxl' if filename.lower().endswith('.xlsx') else 'xlrd'
            xls = pd.ExcelFile(io.BytesIO(file_content), engine=engine)
                
            for s in xls.sheet_names:
                df_raw = pd.read_excel(xls, sheet_name=s, header=None)
                if df_raw.empty: continue

                # 1. Look for Section Titles & Metadata Labels
                for r_idx, row in df_raw.iterrows():
                    row_vals = [str(x) for x in row.tolist() if pd.notna(x) and str(x).strip()]
                    for text in row_vals:
                        if len(text) > 100 or "NMG MARINE" in text.upper(): continue
                        
                        kv = re.search(r'^([^:]{2,50}):\s*(.*)$', text)
                        if kv:
                            data["metadata"][kv.group(1).strip()] = kv.group(2).strip()
                        elif text.isupper() and len(text) > 5 and ":" not in text:
                            # Potential section title inside sheet
                            data["paragraphs"].append(text)

                # 2. Main Table Discovery
                # Strategy: Finding the row with "S.NO" or most non-label headers
                best_header_row = 0
                max_score = -1
                
                for r_idx, row in df_raw.iterrows():
                    cells = [str(c).strip() for c in row if pd.notna(c) and str(c).strip()]
                    if not cells: continue
                    
                    score = 0
                    if any(x in "".join(cells).upper() for x in ["S.NO", "ITEM", "PARTICULARS", "DESCRIPTION"]):
                        score += 50
                    
                    # Count cells that aren't labels (don't end in :)
                    clean_cells = [c for c in cells if not c.endswith(':')]
                    score += len(clean_cells) * 2
                    
                    if score > max_score:
                        max_score = score
                        best_header_row = r_idx

                # Skip metadata rows to read the actual table
                df = pd.read_excel(xls, sheet_name=s, skiprows=best_header_row)
                if not df.empty:
                    table_data = [df.columns.tolist()] + df.values.tolist()
                    cleaned_table = []
                    for r in table_data:
                        # Prune noise rows (section labels like "TODAYS AVERAGES")
                        non_empty = [c for c in r if pd.notna(c) and str(c).strip()]
                        if len(non_empty) == 1 and len(str(non_empty[0])) > 12: continue
                        
                        if not any(pd.notna(x) and str(x).strip() for x in r): continue
                        cleaned_table.append([(str(c) if pd.notna(c) else "") for c in r])
                    
                    if len(cleaned_table) > 1:
                        data["tables"].append(cleaned_table)
            return data
        except Exception as e:
            return {"error": f"Excel Error: {str(e)}"}

class WordParser:
    @staticmethod
    def extract(file_content: bytes):
        try:
            doc = Document(io.BytesIO(file_content))
        except: return {"error": "Invalid Word document"}
        data = {"paragraphs": [], "tables": [], "metadata": {}, "raw_text": ""}
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = []
                seen = set()
                for cell in row.cells:
                    if cell._element not in seen:
                        txt = cell.text.strip()
                        cells.append(txt)
                        kv = re.search(r'^([^:]{2,40}):\s*(.*)$', txt)
                        if kv: data["metadata"][kv.group(1).strip()] = kv.group(2).strip()
                        seen.add(cell._element)
                    else:
                        # For merged cells, maintain grid alignment by adding empty string
                        cells.append("")
                table_rows.append(cells)
            if any(any(r) for r in table_rows):
                data["tables"].append(table_rows)

        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt or "NMG MARINE" in txt.upper(): continue
            data["paragraphs"].append(txt); data["raw_text"] += txt + "\n"
            segs = re.split(r'\t| {3,}', txt)
            for s in segs:
                kv = re.search(r'^([^:]{2,50}):\s*(.*)$', s)
                if kv: data["metadata"][kv.group(1).strip()] = kv.group(2).strip()
        return data

class PDFParser:
    @staticmethod
    def extract(file_content: bytes):
        data = {"paragraphs": [], "tables": [], "metadata": {}, "raw_text": ""}
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        data["raw_text"] += text + "\n"
                        data["paragraphs"].extend([p.strip() for p in text.split('\n') if p.strip()])
                    tables = page.extract_tables()
                    for table in (tables or []):
                        cleaned = [[(str(c).strip() if c else "") for c in r] for r in table]
                        if cleaned: data["tables"].append(cleaned)
        except: pass
        for p in data["paragraphs"]:
            segs = re.split(r'\t| {3,}', p)
            for s in segs:
                kv = re.search(r'^([^:]{2,50}):\s*(.*)$', s)
                if kv: data["metadata"][kv.group(1).strip()] = kv.group(2).strip()
        return data

class SmartParser:
    @staticmethod
    def _extract_text_from_doc_binary(file_content: bytes) -> str:
        """
        Pure-Python fallback: Extract readable text from .doc binary (OLE2 format).
        No external tools needed. Works on any OS.
        """
        import struct
        
        text_parts = []
        
        # Method 1: Try OLE2 WordDocument stream extraction
        try:
            # OLE2 magic number check
            if file_content[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
                # This is a valid OLE2 file, try to find text in the compound doc
                # Look for the FIB (File Information Block) and extract text
                # The text in Word binary is stored as UTF-16LE after specific markers
                
                # Extract UTF-16LE text spans (Word stores text as UTF-16)
                decoded_chunks = []
                i = 0
                while i < len(file_content) - 1:
                    # Look for sequences of printable UTF-16LE characters
                    chunk_start = i
                    chars = []
                    while i < len(file_content) - 1:
                        code = struct.unpack_from('<H', file_content, i)[0]
                        # Printable ASCII range in UTF-16LE, plus common chars
                        if (0x20 <= code <= 0x7E) or code in (0x09, 0x0A, 0x0D, 0x2013, 0x2014, 0x2018, 0x2019, 0x201C, 0x201D):
                            chars.append(chr(code))
                            i += 2
                        else:
                            break
                    
                    if len(chars) >= 4:  # Only keep chunks of 4+ chars
                        decoded_chunks.append(''.join(chars))
                    i += 2
                
                if decoded_chunks:
                    # Find the longest meaningful text blocks
                    long_chunks = [c for c in decoded_chunks if len(c) >= 6]
                    if long_chunks:
                        text_parts.extend(long_chunks)
        except Exception:
            pass
        
        # Method 2: Simple ASCII/Latin1 extraction fallback
        if not text_parts:
            try:
                text = file_content.decode('latin-1', errors='ignore')
                # Remove control characters but keep newlines and tabs
                cleaned = re.sub(r'[^\x09\x0a\x0d\x20-\x7e\xa0-\xff]', ' ', text)
                # Split on large gaps of spaces (indicating field boundaries in Word)
                cleaned = re.sub(r' {4,}', '\n', cleaned)
                cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
                if len(cleaned.strip()) > 20:
                    text_parts = [cleaned.strip()]
            except Exception:
                pass
        
        full_text = '\n'.join(text_parts)
        
        # Clean up noise common in .doc binary extraction
        # Remove very short lines that are likely binary artifacts
        lines = full_text.split('\n')
        clean_lines = []
        for line in lines:
            stripped = line.strip()
            # Keep lines that look like real text (3+ chars, not all special chars)
            if len(stripped) >= 2 and re.search(r'[a-zA-Z0-9]', stripped):
                clean_lines.append(stripped)
        
        return '\n'.join(clean_lines)

    @staticmethod
    def _convert_doc_to_docx(file_content: bytes) -> bytes:
        """Helper to convert .doc to .docx across platforms. Returns None if no conversion tool available."""
        tmp_doc = None
        tmp_docx = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(file_content)
                tmp_doc = tmp.name
            
            tmp_docx = tmp_doc + "x"
            abspath_doc = os.path.abspath(tmp_doc)
            abspath_docx = os.path.abspath(tmp_docx)

            # Prioritize OS-specific best tools
            if os.name == 'nt':
                if win32 is not None:
                    pythoncom.CoInitialize()
                    try:
                        word = win32.Dispatch("Word.Application")
                        word.Visible = False
                        doc = word.Documents.Open(abspath_doc)
                        doc.SaveAs2(abspath_docx, FileFormat=12)
                        doc.Close()
                        word.Quit()
                    finally:
                        pythoncom.CoUninitialize()
            else:
                # Linux/Mac: Try Pandoc first, then LibreOffice
                if shutil.which("pandoc"):
                    subprocess.run(["pandoc", abspath_doc, "-o", abspath_docx], check=False, capture_output=True)
                
                soft = shutil.which("soffice") or shutil.which("libreoffice")
                if not os.path.exists(abspath_docx) and soft:
                    tmp_dir = os.path.dirname(abspath_doc)
                    subprocess.run([soft, "--headless", "--convert-to", "docx", "--outdir", tmp_dir, abspath_doc], check=False, capture_output=True)
                    expected_out = abspath_doc.rsplit('.', 1)[0] + ".docx"
                    if os.path.exists(expected_out) and expected_out != abspath_docx:
                        shutil.move(expected_out, abspath_docx)

            if os.path.exists(abspath_docx):
                with open(abspath_docx, "rb") as f:
                    return f.read()
            
            # No conversion tool available
            return None
        finally:
            if tmp_doc and os.path.exists(tmp_doc):
                try: os.remove(tmp_doc)
                except: pass
            if tmp_docx and os.path.exists(tmp_docx):
                try: os.remove(tmp_docx)
                except: pass

    @staticmethod
    def parse_to_form(file_content: bytes, filename: str):
        fn = filename.lower()
        try:
            if fn.endswith('.doc') and not fn.endswith('.docx'):
                # Try converting to .docx first (best quality)
                docx_content = SmartParser._convert_doc_to_docx(file_content)
                if docx_content:
                    print(f"✅ .doc converted to .docx successfully for: {filename}")
                    raw = WordParser.extract(docx_content)
                else:
                    # Fallback: Pure-Python binary text extraction (no tools needed)
                    print(f"⚠️ No conversion tools available. Using pure-Python .doc extraction for: {filename}")
                    extracted_text = SmartParser._extract_text_from_doc_binary(file_content)
                    if not extracted_text or len(extracted_text) < 10:
                        raise ValueError("Could not extract text from .doc file. Please convert to .docx format for better results.")
                    
                    # Build a raw structure similar to what WordParser returns
                    paragraphs = [line for line in extracted_text.split('\n') if line.strip()]
                    
                    # Try to detect tables (lines with tab-separated values)
                    tables = []
                    regular_paragraphs = []
                    current_table = []
                    for para in paragraphs:
                        if '\t' in para:
                            current_table.append(para.split('\t'))
                        else:
                            if current_table:
                                tables.append(current_table)
                                current_table = []
                            regular_paragraphs.append(para)
                    if current_table:
                        tables.append(current_table)
                    
                    # Detect key:value metadata from paragraphs
                    metadata = {}
                    for para in regular_paragraphs:
                        if ':' in para:
                            parts = para.split(':', 1)
                            if len(parts) == 2 and len(parts[0].strip()) < 50:
                                metadata[parts[0].strip()] = parts[1].strip()
                    
                    raw = {
                        "paragraphs": regular_paragraphs,
                        "tables": tables,
                        "metadata": metadata,
                        "raw_text": extracted_text
                    }

            elif fn.endswith('.docx'): raw = WordParser.extract(file_content)
            elif fn.endswith('.pdf'): raw = PDFParser.extract(file_content)
            elif fn.endswith(('.xlsx', '.xls', '.csv')):
                if fn.endswith('.csv'):
                    df = pd.read_csv(io.BytesIO(file_content))
                    raw = {"paragraphs": [], "tables": [[df.columns.tolist()] + df.values.tolist()], "metadata": {}, "raw_text": ""}
                else: raw = ExcelParser.extract(file_content, filename)
            else:
                 raw = {"error": "Unsupported file format. Please use .docx, .xlsx, or .pdf"}

            if "error" in raw: raise ValueError(raw["error"])
            
            fields = BaseParser.map_to_form_fields(raw)
            if not fields: raise ValueError("Could not find any structure in this document.")
            
            # POPULATE SPREADSHEET MODE: If it's excel, we also provide the raw grid for "Direct Link" mode
            spreadsheet_data = None
            if fn.endswith(('.xlsx', '.xls')) and raw.get("tables"):
                main_table = raw["tables"][0]
                if len(main_table) > 0:
                    cols = [chr(65 + i) for i in range(len(main_table[0]))] # A, B, C...
                    spreadsheet_data = {
                        "rows": main_table,
                        "columnNames": cols
                    }

            return {
                "title": filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').title(),
                "fields": fields,
                "spreadsheet_data": spreadsheet_data
            }
        except Exception as e:
            raise ValueError(str(e))
